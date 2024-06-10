from flask import Flask, request, jsonify, render_template, session
from flask_session import Session
import asyncio
from PyPDF2 import PdfReader
from docx import Document
import docx2txt
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
import json
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores.faiss import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain_core.messages import AIMessage, HumanMessage
from langchain_core.output_parsers import StrOutputParser
import requests
from bs4 import BeautifulSoup
from werkzeug.utils import secure_filename
import pandas as pd
import spacy
import requests
from bs4 import BeautifulSoup
import PyPDF2
from neo4j import GraphDatabase
from langchain_core.runnables import (
    RunnableBranch,
    RunnableLambda,
    RunnableParallel,
    RunnablePassthrough,
)
from langchain_core.prompts import ChatPromptTemplate
from langchain.graphs import Neo4jGraph
from langchain_core.prompts.prompt import PromptTemplate
from langchain_core.pydantic_v1 import BaseModel, Field
from typing import Tuple, List, Optional
from langchain_core.messages import AIMessage, HumanMessage
from langchain_core.output_parsers import StrOutputParser
import os
from langchain_community.graphs import Neo4jGraph
from langchain_community.llms import Ollama
from langchain.text_splitter import TokenTextSplitter
from langchain_experimental.graph_transformers import LLMGraphTransformer
from neo4j import GraphDatabase
from langchain_community.vectorstores import Neo4jVector

from langchain_community.vectorstores.neo4j_vector import remove_lucene_chars
from langchain_core.runnables import ConfigurableField, RunnableParallel, RunnablePassthrough
from langchain.docstore.document import Document 
'''part of the LangChain library and is used to represent a piece of text or document data along with its metadata. 
It is a common data structure used throughout the LangChain library for storing and processing text data.
In your code, you are using the Document class to create new Document objects from the preprocessed text data. 
Specifically, you are creating a dictionary with the page_content (the actual text content) and metadata (additional information about the text, such as the source), 
and then creating a Document object using that dictionary.'''
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.embeddings import HuggingFaceEmbeddings
from py2neo import Graph
from langchain_community.vectorstores import Neo4jVector
from langchain.embeddings import OpenAIEmbeddings
import os
from langchain_experimental.llms.ollama_functions import OllamaFunctions
from langchain_community.llms import Ollama
from langchain_community.chat_models import ChatOllama
from langchain.prompts import PromptTemplate
from langchain.output_parsers import StructuredOutputParser
from pydantic import BaseModel
from typing import List

app = Flask(__name__)
app.secret_key = 'supersecretkey'
app.config['SESSION_TYPE'] = 'filesystem'
Session(app)
URI = "bolt://localhost:7687"
graph = Neo4jGraph(url=URI, username="neo4j", password="password")

# Ensure the environment variables are loaded
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Function to handle asyncio event loop
def get_or_create_eventloop():
    try:
        return asyncio.get_event_loop()
    except RuntimeError as ex:
        if "There is no current event loop in thread" in str(ex):
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            return asyncio.get_event_loop()
get_or_create_eventloop()

def get_text_from_doc(doc_file):
    document = Document(doc_file)
    return "\n".join([paragraph.text for paragraph in document.paragraphs])

def get_text_from_docx(docx_file):
    # Save the uploaded file to a temporary location to be processed by docx2txt
    temp_file_path = "temp.docx"
    with open(temp_file_path, "wb") as f:
        f.write(docx_file.read())
    return docx2txt.process(temp_file_path)

def get_text_from_txt(txt_file):
    return txt_file.read().decode("utf-8")

def get_text_from_pdf(pdf_file):
    text = ""
    pdf_reader = PdfReader(pdf_file)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def get_text_from_files(files):
    text = ""
    for file in files:
        if file.filename.endswith(".pdf"):
            text += get_text_from_pdf(file)
        elif file.filename.endswith(".doc"):
            text += get_text_from_doc(file)
        elif file.filename.endswith(".docx"):
            text += get_text_from_docx(file)
        elif file.filename.endswith(".txt"):
            text += get_text_from_txt(file)
        else:
            return f"Unsupported file type: {file.filename}"
    return text

def get_text_from_url(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, "html.parser")
        text = ' '.join(p.get_text() for p in soup.find_all('p'))
        return text
    except Exception as e:
        return f"Error fetching the URL: {e}"

def get_text_chunks(text):
    
    text_splitter = TokenTextSplitter(chunk_size=512, chunk_overlap=24)
    documents = text_splitter.split_text(text[:3])
    
    return documents
# take foldername input 
def get_vector_store(text_chunks, usersession):

# Preprocess the documents to convert lists to tuples

    preprocessed_documents = []
    for doc in text_chunks:
        data = {
        "page_content": doc.page_content,
        "metadata": {
            "source": tuple(doc.metadata["source"]) if isinstance(doc.metadata["source"], list) else doc.metadata["source"]
        }
    }
        preprocessed_documents.append(Document(**data))
    llm = OllamaFunctions(model="llama3")
    llm_transformer = LLMGraphTransformer(llm=llm)
    graph_documents = llm_transformer.convert_to_graph_documents(text_chunks)
    graph.add_graph_documents(
    graph_documents,
    baseEntityLabel=True,
    include_source=True
    )

    from langchain_community.embeddings import OllamaEmbeddings


    embeddings = OllamaEmbeddings(model="mxbai-embed-large")
    session['vector_index'] = Neo4jVector.from_existing_graph(
    embeddings,
    url="bolt://localhost:7687",
    username='neo4j',
    password='password',
    search_type="hybrid",
    node_label="Document",
    text_node_properties=["text"],
    embedding_node_property="embedding"
)
# Retriever
    from py2neo import                        Graph
    graph=Graph('bolt://localhost:7687', name='neo4j')

    graph.run(
    "CREATE FULLTEXT INDEX entity IF NOT EXISTS FOR (e:__Entity__) ON EACH [e.id]"
)


def get_conversational_chain():
    from langchain_community.chat_models import ChatOllama

    llm = ChatOllama(model="llama3", format="json", temperature=0)
    from langchain.chains import RetrievalQA

    session['vector_qa'] = RetrievalQA.from_chain_type(
    llm,
    chain_type="stuff",
    retriever=session['vector_index'].as_retriever()
)
  
    return session['vector_qa']

def user_input(user_question):
    response = session['vector_qa'].run(user_question)
    return response

@app.route('/', methods=['GET', 'POST'])
def index():
    message = None  # Initialize a message variable
    file_details = []  # Initialize a list to store file details
    url_displayed = session.get('url_input', '')  # Retrieve the stored URL or set it to empty string

    if 'session_id' not in session:
        session['session_id'] = os.urandom(24).hex()
        session['chat_history'] = [
            AIMessage(content="Hello! I'm a document assistant. Ask me anything about the documents you upload."),
        ]

    if request.method == 'POST':
        files = request.files.getlist("files")
        url_input = request.form.get("url_input")
        raw_text = ""
        session["input_language"] = int(request.form.get("input_language"))
        
        session["output_language"] = int(request.form.get("output_language"))
        # Process files
        if files and files[0].filename != '':
            valid_files = all(f.filename.endswith(('.pdf', '.doc', '.docx', '.txt')) for f in files)
            if valid_files:
                raw_text += get_text_from_files(files)
                message = "Files successfully uploaded."

                # Get file details for display
                for file in files:
                    file_details.append({"name": file.filename})
            else:
                message = "Please upload files in PDF, DOC, DOCX, or TXT format."

        # Process URL
        if url_input:
            url_text = get_text_from_url(url_input)
            raw_text += " " + url_text  # Concatenate URL text with existing text
             # Debug print to check what is being added
            message = "Files and URL processed successfully. URL : "+ url_input

            session['url_input'] = url_input  # Store the URL in the session


        if raw_text:
            text_chunks = get_text_chunks(raw_text)
            get_vector_store(text_chunks, session['session_id'])

    chat_history = session.get('chat_history', [])
    return render_template('index.html', chat_history=chat_history, message=message, file_details=file_details, url_displayed=url_displayed)


@app.route('/ask', methods=['POST'])
def ask():
    user_query = request.json.get("question")
    if user_query and user_query.strip():
        session['chat_history'].append(HumanMessage(content=user_query))
        response = user_input(user_query, session['session_id'])

        res = response["output_text"]
        session['chat_history'].append(AIMessage(content=res))
        print(request.form.get("input_language"))
        

        if int(session["output_language"]) != 23:
            payload = {
        "source_language": session["input_language"],
        "content": res,
        "target_language": session["output_language"]
      }
            res = json.loads(requests.post('http://127.0.0.1:8000/scaler/translate', json=payload).content)
            res = res['translated_content']

        return jsonify({"answer": res, "url": session.get('url_input', '')})
    return jsonify({"error": "Invalid input"})

if __name__ == "__main__":
    app.run(debug=True)